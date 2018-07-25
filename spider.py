from bs4 import BeautifulSoup
import requests
from selenium import webdriver
import time
from PIL import Image
import os
from docx import Document
from docx.shared import Inches


class spider_image:
    browser = webdriver.Chrome()

    def get_data_from_website(self):
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/67.0.3396.99 Safari/537.36',
            'Cookie': 'Hm_lvt_a1c802481a8c3cc5ae9491bb48e30192=1531991086,1532268251,1532405598;'
                      ' PHPSESSID=ST-305156-gBocuNO9zxU9iIJ1gch1-castwo; '
                      'Hm_lpvt_a1c802481a8c3cc5ae9491bb48e30192=1532420693'
        }
        url = 'http://zjc.ncu.edu.cn/jy/index.php?c=channel&a=type&tid=68'
        response = requests.get(url, headers=headers)
        response.encoding = 'utf-8'
        html = response.text
        soup = BeautifulSoup(html, 'lxml')
        infoes = soup.select('body > div > div.am-u-md-9 > ul.am-list > li > a ')
        DATA = []
        for info in infoes:
            href = 'http://zjc.ncu.edu.cn' + info.get('href')
            info_text = info.getText()
            data = {
                'href': href,
                # 单位名称
                'name': info_text[14:],
                # 发布时间
                'pub_time': info_text[0:13]
            }
            DATA.append(data)
        return DATA

    def jump_to(self, url):
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) '
                          'AppleWebKit/537.36 (KHTML, like Gecko) Chrome/66.0.3359.181 Safari/537.36',
            'Cookie': 'Hm_lvt_a1c802481a8c3cc5ae9491bb48e30192=1531814547,1531814551,1531873009,1531968482; '
                      'PHPSESSID=ST-298502-2dSpTrejziljjCGV7cFL-castwo; Hm_lpvt_a1c802481a8c3cc5ae9491bb48e30192=1531968834',
            'Connection': 'keep - alive'
        }
        browser = self.browser
        browser.get(url)
        # 如果没有跳转到目标页面，可能是需要登录
        if not browser.current_url == url:
            try:
                browser.find_element_by_id('username').send_keys('6103215038')
                browser.find_element_by_id('password').send_keys('300071')
                browser.find_element_by_css_selector('#fm1 > input[type="image"]:nth-of-type(1)').click()
                # 这里注意要等待2秒 否则不会正常跳转
                time.sleep(1)
                relocation_str = 'window.location.href = "' + url + '";'
                browser.execute_script(relocation_str)
                time.sleep(1)
            except Exception:
                print('session error')

    def capture(self, pix_h, filename):
        """chrome截屏
        pix_h- 窗口高
        filename-生成截图的文件名
        """
        try:
            # 全屏
            self.browser.maximize_window()
            time.sleep(2)
            img_list = []
            i = 0
            while i < 20:
                # 滚屏
                js = "var q=document.documentElement.scrollTop=" + str(i * pix_h) + ";"
                self.browser.execute_script(js)
                js1 = "return document.documentElement.scrollHeight.toString()+','+document.documentElement.scrollTop.toString()"
                js1_result = self.browser.execute_script(js1)
                real_scroll_h, real_top = js1_result.split(',')[0], round(float(js1_result.split(',')[1]))
                # real_scroll_h, real_top 是当前滚动条长度和当前滚动条的top，作为是否继续执行的依据，由于存在滚动条向下拉动后会加载新内容的情况，所以需要以下的判断
                # 如果这次设置的top成功，则继续滚屏
                print('real_top = {}, ({}*pix_h) = {}'.format(real_top, i, i*pix_h))
                if str(real_top) == str(i * pix_h):
                    i += 1
                    self.browser.save_screenshot(filename + str(i) + '.png')
                    img_list.append(filename + str(i) + '.png')
                    last_t = real_top
                    print('real_top={}, last_t={}'.format(real_top, last_t))
                else:
                    # 如果本次设置失败，看这次的top和上一次记录的top值是否相等，相等则说明没有新加载内容，且已到页面底，跳出循环
                    if real_top != last_t:
                        print('real_top={}, last_t={}'.format(real_top, last_t))
                        last_t = real_top
                    else:
                        print('real_top={}, last_t={}'.format(real_top, last_t))
                        self.browser.save_screenshot(filename + str(i + 1) + '.png')
                        img_list.append(filename + str(i + 1) + '.png')
                        break
            return self.image_merge(img_list, 'table.png')
        except Exception as e:
            print(e)

    def image_merge(self, images, output_name='merge.png'):
        """垂直合并多张图片
        images - 要合并的图片路径列表
        ouput_dir - 输出路径
        output_name - 输出文件名
        restriction_max_width - 限制合并后的图片最大宽度，如果超过将等比缩小
        restriction_max_height - 限制合并后的图片最大高度，如果超过将等比缩小
        """
        max_width = 0
        total_height = 0
        # 计算合成后图片的宽度（以最宽的为准）和高度
        for img_path in images:
            if os.path.exists(img_path):
                img = Image.open(img_path)
                width, height = img.size
                if width > max_width:
                    max_width = width
                total_height += height
        # 产生一张空白图
        new_img = Image.new('RGB', (max_width, total_height), 255)
        # 合并
        x = y = 0
        for img_path in images:
            if os.path.exists(img_path):
                img = Image.open(img_path)
                width, height = img.size
                new_img.paste(img, (x, y))
                y += height

        save_path = output_name
        print(save_path)
        # 剪裁table
        table = self.browser.find_element_by_tag_name('table')
        left = table.location['x']
        top = table.location['y'] + 100
        right = table.location['x'] + table.size['width'] + 400
        bottom = table.location['y'] + table.size['height'] + 600
        new_img = new_img.crop((left, top, right, bottom))
        new_img.save(save_path)

        for img_path in images:
            os.remove(img_path)
        return new_img

    def is_table_exist(self, tag_name='table'):
        flag = True
        browser = self.browser
        try:
            browser.find_element_by_tag_name(tag_name)
            return flag
        except:
            flag = False
            return flag

    def save_docx(self, url):
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/67.0.3396.99 Safari/537.36',
            'Cookie': 'Hm_lvt_a1c802481a8c3cc5ae9491bb48e30192=1531991086,1532268251,1532405598; '
                      'PHPSESSID=ST-305156-gBocuNO9zxU9iIJ1gch1-castwo; '
                      'Hm_lpvt_a1c802481a8c3cc5ae9491bb48e30192=1532420675'
        }
        response = requests.get(url, headers=headers)
        response.encoding = 'utf-8'
        html = response.text
        soup = BeautifulSoup(html, 'lxml')
        time = soup.select('body > div > div.am-u-md-9 > div.defined-zp-content > span:nth-of-type(1)')
        text = soup.select('body > div > div.am-u-md-9 > div.defined-zp-content > span:nth-of-type(2) > p')
        # title = soup.select('body > div > div.am-u-md-9 > div.defined-zp-content > span:nth-of-type(2) > h1 > span')
        name = soup.select('body > div > div.am-u-md-9 > div.defined-zp-content > section:nth-of-type(1) > h2 > span:nth-of-type(3) > strong')
        document = Document()
        # document.add_heading(title[0].text.strip(), level=0)
        document.add_heading(time[0].text.strip(), level=1)
        for i in range(len(text)):
            if len(text[i].text) != 0:
                document.add_paragraph(text[i].get_text().strip())

        # 添加截图
        if self.is_table_exist():
            document.add_paragraph('表格见下方,具体请见网页：{}'.format(url))
            self.capture(spider_image.browser.get_window_size()['height'], url[-4:])
            document.add_picture(r'table.png', width=Inches(5.5))
        else:
            document.add_paragraph('具体请见网页：{}'.format(url))

        file_name = name[0].get_text().replace(",", "").replace('"', '') + '.docx'
        document.save(file_name)
        if os.path.exists(r'table.png'):
            os.remove(r'table.png')

    def save(self, data):
        for d in data:
            url = d['href']
            self.jump_to(url)
            self.save_docx(url)


if __name__ == '__main__':
    spider_image = spider_image()
    DATA = spider_image.get_data_from_website()
    spider_image.save(DATA)
